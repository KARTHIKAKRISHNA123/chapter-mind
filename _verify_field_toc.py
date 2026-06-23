# -*- coding: utf-8 -*-
import os, sys
sys.path.insert(0, os.path.dirname(__file__))
from book_splitter.adapters.docx_toc import resolve_field_toc

# ---- (A) REAL data extracted from the uploaded pandoc Tamil book -----------
T = "வானம் தொட்டு விடும் தூரம்தான்"
field_titles = [T, "அச்சுப்பதிப்பு", "சமர்ப்பணம்", "அணிந்துரை", "பதிப்புரை",
                "என்னுரை", "தரிசு நிலம்", "அந்திம கிரியை", T, "அம்மா",
                "பிணைப்பால் தொடரும் உறவுகள்", "இருட்டில் ஒரு வெளிச்சம்",
                "இதுவும் கடந்து போகும்", "அப்பா", "கொஞ்சம் கல்வி நிறைய காசு",
                "வேலை", "வீழ்வேன் என்று நினைத்தாயோ ?", "அன்னபூரணி",
                "நான் செய்தது", "கணியம் அறக்கட்டளை"]
# (index, text, style_level)  — H1 -> level 1, H2 -> level 2 (pandoc off-by-one)
heading_pool = [
    (1, "Table of Contents", 1), (23, T, 1), (29, T, 1), (55, T, 1),
    (78, "அச்சுப்பதிப்பு", 2), (81, "சமர்ப்பணம்", 2), (85, "அணிந்துரை", 2),
    (109, "பதிப்புரை", 2), (128, "என்னுரை", 2), (143, "தரிசு நிலம்", 1),
    (208, "அந்திம கிரியை", 1), (268, T, 1), (415, "அம்மா", 1),
    (489, "பிணைப்பால் தொடரும் உறவுகள்", 1), (604, "இருட்டில் ஒரு வெளிச்சம்", 1),
    (713, "இதுவும் கடந்து போகும்", 1), (790, "அப்பா", 1),
    (884, "கொஞ்சம் கல்வி நிறைய காசு", 1), (973, "வேலை", 1),
    (1053, "வீழ்வேன் என்று நினைத்தாயோ ?", 1), (1150, "அன்னபூரணி", 1),
    (1292, "நான் செய்தது", 1), (1336, "கணியம் அறக்கட்டளை", 2),
]
res = resolve_field_toc(field_titles, heading_pool)
idxs = [e["index"] for e in res] if res else None
expected = [23, 78, 81, 85, 109, 128, 143, 208, 268, 415, 489, 604, 713, 790,
            884, 973, 1053, 1150, 1292, 1336]
print("(A) real-book resolver: entries =", len(res) if res else 0)
print("    indices match expected:", idxs == expected)
if idxs != expected:
    print("    got:", idxs)

# ---- (B) synthetic end-to-end: real TOC field, no _Toc bookmarks -----------
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _fld(p, t):
    r = p.add_run(); e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), t); r._r.append(e)
def _instr(p, txt):
    r = p.add_run(); e = OxmlElement('w:instrText'); e.set(qn('xml:space'), 'preserve'); e.text = txt; r._r.append(e)

d = Document()
d.add_heading("Table of Contents", level=1)
fp = d.add_paragraph(); _fld(fp, 'begin'); _instr(fp, ' TOC \\o "1-3" '); _fld(fp, 'separate')
for t in ["Chapter One", "Chapter Two", "Chapter Three", "About the Author"]:
    d.add_paragraph(t)
ep = d.add_paragraph(); _fld(ep, 'end')
# body (note: title-page repeat of "Chapter One" before the real chapters)
d.add_heading("Chapter One", level=1); d.add_paragraph("Body text for the title page repeat here, long enough to be prose.")
for t in ["Chapter One", "Chapter Two", "Chapter Three"]:
    d.add_heading(t, level=1)
    d.add_paragraph("This is body content for the chapter, written to be clearly prose.")
d.add_heading("About the Author", level=2)
d.add_paragraph("Biographical note that is also ordinary prose content here.")
out = os.path.join(os.path.dirname(__file__), "tests", "golden", "_tmp_field_toc.docx")
d.save(out)

from book_splitter.adapters.registry import get_adapter
from book_splitter.detector import detect
doc = get_adapter(out).load(out)
print("\n(B) synthetic end-to-end:")
print("    toc_entries:", None if doc.toc_entries is None else len(doc.toc_entries))
plan = detect(doc, level_filter="auto")
print("    abstained:", plan.abstained, "chapters:", len(plan.chapters),
      "body:", len(plan.body_chapters))
for c in plan.chapters:
    print(f"       {c.role:<12} {c.title}")
os.remove(out)
